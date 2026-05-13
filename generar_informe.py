from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

# ---- PORTADA ----
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("INFORME TÉCNICO DE DESPLIEGUE")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Intercagram en Microsoft Azure")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 102, 153)

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run(f"Fecha: {datetime.date.today().strftime('%d/%m/%Y')}\n").font.size = Pt(12)
info.add_run("Plataforma: Microsoft Azure\n").font.size = Pt(12)
info.add_run("Máquina Virtual: RPOO\n").font.size = Pt(12)
info.add_run("IP Pública: 20.151.96.66").font.size = Pt(12)

doc.add_page_break()

# ---- 1. INTRODUCCIÓN ----
doc.add_heading("1. Introducción", level=1)
doc.add_paragraph(
    "Intercagram es una aplicación de red social inspirada en Instagram, desarrollada como proyecto "
    "académico. Permite a los usuarios registrarse, publicar imágenes, dar likes, comentar, y gestionar "
    "amistades."
)
doc.add_paragraph(
    "El presente informe documenta el proceso de despliegue de la aplicación en Microsoft Azure, "
    "utilizando una Máquina Virtual (VM) con sistema operativo Ubuntu Server 24.04 LTS."
)

# ---- 2. ARQUITECTURA ----
doc.add_heading("2. Arquitectura del Proyecto", level=1)

table = doc.add_table(rows=6, cols=3)
table.style = 'Light Shading Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Componente', 'Tecnología', 'Versión']
data = [
    ['Backend API', 'Laravel (PHP)', '12.x / PHP 8.3'],
    ['Frontend Web', 'Ionic + Angular', 'Ionic 8 / Angular 20'],
    ['Base de Datos', 'MySQL', '8.0'],
    ['Autenticación', 'Laravel Sanctum (tokens)', '4.x'],
    ['Servidor Web', 'Apache', '2.4.58'],
]

for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph()
doc.add_paragraph(
    "La aplicación sigue una arquitectura cliente-servidor: el frontend Ionic/Angular se comunica "
    "con el backend Laravel a través de una API REST, utilizando tokens de autenticación (Sanctum). "
    "Las imágenes se almacenan en el sistema de archivos local de la VM."
)

# ---- 3. INFRAESTRUCTURA AZURE ----
doc.add_heading("3. Infraestructura en Microsoft Azure", level=1)

doc.add_heading("3.1 Máquina Virtual", level=2)
table2 = doc.add_table(rows=6, cols=2)
table2.style = 'Light Shading Accent 1'

vm_data = [
    ('Nombre', 'RPOO'),
    ('Sistema Operativo', 'Ubuntu Server 24.04 LTS'),
    ('IP Pública', '20.151.96.66'),
    ('Grupo de Recursos', 'RPOO_group_04301447'),
    ('Suscripción', '25a0f80b-8d95-4aab-af55-68e5bb200983'),
]

for i, (k, v) in enumerate(vm_data):
    table2.rows[i].cells[0].text = k
    table2.rows[i].cells[1].text = v
    table2.rows[i].cells[0].paragraphs[0].runs[0].bold = True if table2.rows[i].cells[0].paragraphs[0].runs else False

doc.add_paragraph()

doc.add_heading("3.2 Seguridad de Red (NSG)", level=2)
doc.add_paragraph(
    "Se configuraron las siguientes reglas de seguridad en el Grupo de Seguridad de Red (NSG):"
)
items = [
    "Puerto 22 (SSH): acceso solo con clave privada (.pem)",
    "Puerto 80 (HTTP): redirige a HTTPS",
    "Puerto 443 (HTTPS): abierto al público para la aplicación web con SSL",
    "Puerto 3306 (MySQL): cerrado (solo acceso local desde la VM)",
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("3.3 Configuración de HTTPS", level=2)
doc.add_paragraph(
    "Se generó un certificado SSL autofirmado y se configuró Apache para servir "
    "la aplicación exclusivamente sobre HTTPS en el puerto 443:"
)
doc.add_paragraph(
    "sudo openssl req -x509 -nodes -days 365 -newkey rsa:2048 \\\n"
    "-keyout /etc/ssl/private/intercagram.key \\\n"
    "-out /etc/ssl/certs/intercagram.crt \\\n"
    "-subj '/CN=20.151.96.66'", style='Quote'
)
doc.add_paragraph(
    "Se creó el VirtualHost intercagram-ssl.conf en Apache con SSLEngine on, "
    "apuntando a /var/www/html/intercagram/public como DocumentRoot. "
    "El tráfico HTTP (puerto 80) se redirige automáticamente a HTTPS mediante Redirect permanent."
)

# ---- 4. PROCESO DE DESPLIEGUE ----
doc.add_heading("4. Proceso de Despliegue", level=1)

doc.add_heading("4.1 Repositorio GitHub", level=2)
doc.add_paragraph(
    "El código fuente se alojó en GitHub: https://github.com/HYNR24/intercagram"
)
doc.add_paragraph(
    "Esto permitió clonar el repositorio directamente en la VM de Azure para su despliegue."
)

doc.add_heading("4.2 Instalación del Stack LAMP", level=2)
doc.add_paragraph(
    "Se instaló el stack LAMP (Linux, Apache, MySQL, PHP) en la VM de Azure mediante los siguientes comandos:"
)
doc.add_paragraph("sudo apt install -y apache2 mysql-server php8.3 php8.3-cli php8.3-mbstring php8.3-xml php8.3-curl php8.3-mysql php8.3-zip php8.3-gd php8.3-bcmath composer git unzip", style='Quote')

doc.add_heading("4.3 Configuración de la Base de Datos", level=2)
doc.add_paragraph(
    "Se creó una base de datos MySQL llamada 'intercagram_db' con un usuario dedicado:"
)
doc.add_paragraph("CREATE DATABASE intercagram_db;\nCREATE USER 'intercagram'@'localhost' IDENTIFIED BY '********';\nGRANT ALL PRIVILEGES ON intercagram_db.* TO 'intercagram'@'localhost';", style='Quote')

doc.add_heading("4.4 Despliegue del Backend (Laravel)", level=2)
doc.add_paragraph(
    "Se clonó el repositorio en /var/www/html/intercagram/ y se configuró Laravel:"
)
steps = [
    "Configuración del archivo .env con conexión a MySQL",
    "Instalación de dependencias con Composer",
    "Generación de APP_KEY",
    "Ejecución de migraciones (php artisan migrate)",
    "Creación del enlace simbólico de almacenamiento (php artisan storage:link)",
    "Configuración de VirtualHost en Apache",
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading("4.5 Despliegue del Frontend (Ionic/Angular)", level=2)
doc.add_paragraph(
    "El frontend se compiló localmente con Angular y se desplegó en la VM en la ruta /app/:"
)
steps2 = [
    "Configuración de la URL del API en api.ts y auth.ts (http://20.151.96.66/api/)",
    "Configuración de la URL de imágenes en feed.page.ts y friends.page.ts",
    "Compilación con ng build --configuration production",
    "Configuración de baseHref en /app/ en angular.json",
    "Subida de la carpeta www/ a la VM",
    "Configuración de .htaccess para enrutamiento Angular",
]
for s in steps2:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading("4.6 Configuración de HTTPS en Apache", level=2)
doc.add_paragraph(
    "Se implementó HTTPS en la aplicación mediante un certificado SSL autofirmado:"
)
steps3 = [
    "Generación de certificado SSL autofirmado con OpenSSL (RSA 2048, validez 365 días)",
    "Creación del VirtualHost intercagram-ssl.conf en Apache con SSLEngine on",
    "Configuración del DocumentRoot en /var/www/html/intercagram/public",
    "Apertura del puerto 443 en el NSG de Azure",
    "Redirección automática de HTTP a HTTPS mediante Redirect permanent",
    "Actualización de APP_URL en .env del backend a https://20.151.96.66",
]
for s in steps3:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading("4.7 Implementación de Cámara WebRTC", level=2)
doc.add_paragraph(
    "Se implementó la captura de fotos directamente desde la cámara del dispositivo "
    "utilizando la API WebRTC (getUserMedia), reemplazando el <input capture='environment'> "
    "que solo funcionaba en dispositivos móviles:"
)
steps4 = [
    "Creación de vista previa en vivo (video tag) con stream de getUserMedia",
    "Botón de captura que toma una foto del video y la dibuja en un canvas",
    "Botón de volteo de cámara (facingMode: 'user' ↔ 'environment') para cámara frontal/posterior",
    "Efecto espejo en la vista previa de la cámara frontal mediante transform: scaleX(-1)",
    "Opción alternativa 'Elegir archivo' con <input type='file'> para subir imágenes existentes",
    "Guardia de doble clic: bandera uploading que deshabilita el botón y muestra 'Publicando...'",
    "Toast de confirmación verde al crear publicación ('Publicación creada', 1.5 segundos)",
    "Navegación automática al feed tras publicación exitosa",
]
for s in steps4:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading("4.8 Corrección de Error 500 en Autenticación", level=2)
doc.add_paragraph(
    "Se solucionó un error crítico que causaba una respuesta HTTP 500 en lugar de 401 "
    "cuando un usuario no autenticado intentaba acceder a rutas protegidas de la API:"
)
doc.add_paragraph(
    "La causa era que Laravel intentaba redirigir a una ruta llamada 'login' (propia de "
    "aplicaciones web con sesiones), inexistente en una API stateless. Se agregó un callback "
    "en bootstrap/app.php que retorna null para rutas API, permitiendo que la excepción "
    "AuthenticationException se maneje correctamente con un response JSON 401:"
)
doc.add_paragraph(
    "->withRouting(\n"
    "    api: __DIR__.'/../routes/api.php',\n"
    "    then: function () {\n"
    "        RedirectIfAuthenticated::redirectUsing(fn () => null);\n"
    "    },\n"
    ")", style='Quote'
)
doc.add_paragraph(
    "Además se agregó un renderizador personalizado para AuthenticationException que "
    "retorna status 401 con mensaje 'Unauthenticated' en formato JSON."
)

doc.add_heading("4.9 Migración de URLs de HTTP a HTTPS", level=2)
doc.add_paragraph(
    "Todas las URLs del frontend Ionic/Angular se actualizaron de HTTP a HTTPS para "
    "evitar errores de contenido mixto (Mixed Content) y permitir el acceso a la cámara "
    "(getUserMedia requiere contexto seguro - HTTPS):"
)
steps5 = [
    "api.ts: baseURL cambiado de http://20.151.96.66/api/ a https://20.151.96.66/api/",
    "auth.ts: URL de autenticación actualizada a https://20.151.96.66/api/",
    "feed.page.ts: imageBaseUrl cambiado a https://20.151.96.66/storage/",
    "friends.page.ts: imageBaseUrl cambiado a https://20.151.96.66/storage/",
    ".env del backend: APP_URL actualizado a https://20.151.96.66",
]
for s in steps5:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading("4.10 Limpieza de Publicaciones Duplicadas", level=2)
doc.add_paragraph(
    "Se identificaron y eliminaron 4 publicaciones duplicadas (IDs 13-16) en la base de datos "
    "MySQL, que fueron generadas durante pruebas de la funcionalidad de cámara y carga de imágenes."
)

# ---- 5. ENDPOINTS API ----
doc.add_heading("5. Endpoints de la API", level=1)
doc.add_paragraph(
    "La API REST expone los siguientes endpoints, todos accesibles desde https://20.151.96.66/api/:"
)

endpoints = [
    ('Autenticación', [
        ('POST /api/register', 'Registro de nuevo usuario'),
        ('POST /api/login', 'Inicio de sesión'),
        ('POST /api/logout', 'Cerrar sesión (token requerido)'),
        ('GET /api/me', 'Perfil del usuario autenticado'),
    ]),
    ('Publicaciones', [
        ('GET /api/posts', 'Listar publicaciones del feed'),
        ('POST /api/posts', 'Crear publicación (imagen + caption)'),
        ('GET /api/posts/{id}', 'Ver publicación'),
        ('DELETE /api/posts/{id}', 'Eliminar publicación'),
        ('POST /api/posts/{id}/like', 'Dar like'),
        ('DELETE /api/posts/{id}/like', 'Quitar like'),
    ]),
    ('Comentarios', [
        ('GET /api/posts/{id}/comments', 'Ver comentarios'),
        ('POST /api/posts/{id}/comments', 'Comentar'),
    ]),
    ('Amistades', [
        ('GET /api/friends', 'Listar amigos'),
        ('POST /api/users/{username}/friend', 'Enviar solicitud'),
        ('DELETE /api/users/{username}/friend', 'Eliminar amigo'),
        ('POST /api/friendships/{id}/accept', 'Aceptar solicitud'),
        ('GET /api/friendships/pending', 'Solicitudes pendientes'),
        ('GET /api/users/suggested', 'Usuarios sugeridos'),
    ]),
    ('Usuarios', [
        ('GET /api/users/search?q=', 'Buscar usuarios'),
        ('GET /api/users/{username}/posts', 'Publicaciones de usuario'),
    ]),
]

for category, endpoints_list in endpoints:
    doc.add_heading(category, level=2)
    table3 = doc.add_table(rows=len(endpoints_list) + 1, cols=2)
    table3.style = 'Light Shading Accent 1'
    table3.rows[0].cells[0].text = 'Endpoint'
    table3.rows[0].cells[1].text = 'Descripción'
    for i, (ep, desc) in enumerate(endpoints_list):
        table3.rows[i + 1].cells[0].text = ep
        table3.rows[i + 1].cells[1].text = desc
    doc.add_paragraph()

# ---- 6. PRUEBAS ----
doc.add_heading("6. Pruebas y Resultados", level=1)
doc.add_paragraph(
    "Se realizaron pruebas de funcionamiento con los siguientes resultados:"
)

test_data = [
    ('Prueba', 'Resultado'),
    ('POST /api/register', 'Crea usuario y devuelve token'),
    ('POST /api/login', 'Autentica y devuelve token'),
    ('GET /api', 'JSON con documentación de endpoints'),
    ('GET /app/', 'Frontend Ionic carga correctamente sobre HTTPS'),
    ('GET /app/login', 'Ruteo Angular funciona'),
    ('GET /storage/posts/*', 'Imágenes se sirven correctamente'),
    ('GET https://20.151.96.66/api/me (sin token)', '401 Unauthenticated (JSON)'),
    ('GET https://20.151.96.66/api/me (con token)', '200 OK, datos del usuario'),
    ('GET https://20.151.96.66/api/posts', '200 OK, feed de publicaciones'),
    ('POST https://20.151.96.66/api/posts', '201 Created con imagen'),
    ('GET https://20.151.96.66/api/friends', '200 OK, lista de amigos'),
    ('Cámara WebRTC (getUserMedia)', 'Stream de video en vivo, captura y carga'),
    ('Doble clic en Publicar', 'Segundo clic bloqueado, muestra "Publicando..."'),
    ('Toast de confirmación', 'Mensaje verde "Publicación creada" 1.5s, navega a /feed'),
    ('Redirección HTTP→HTTPS', 'http://20.151.96.66 redirige a https://20.151.96.66'),
]

table4 = doc.add_table(rows=len(test_data), cols=2)
table4.style = 'Light Shading Accent 1'
for i, (t, r) in enumerate(test_data):
    table4.rows[i].cells[0].text = t
    table4.rows[i].cells[1].text = r

doc.add_paragraph()
doc.add_paragraph(
    "Ejemplo de respuesta de la API (registro de usuario):"
)

example_json = '''{
    "user": {
        "name": "test",
        "email": "test@test.com",
        "profile": {
            "username": "test123"
        }
    },
    "token": "1|xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx"
}'''
doc.add_paragraph(example_json, style='Quote')

# ---- 7. CONCLUSIONES ----
doc.add_heading("7. Conclusiones", level=1)
doc.add_paragraph(
    "1. El proyecto Intercagram fue desplegado exitosamente en Microsoft Azure, "
    "quedando accesible desde cualquier dispositivo con conexión a internet."
)
doc.add_paragraph(
    "2. La combinación Laravel (backend) + Ionic/Angular (frontend) funciona correctamente "
    "en un entorno de producción con Apache y MySQL."
)
doc.add_paragraph(
    "3. Azure VM proporciona un entorno flexible y escalable para aplicaciones web, "
    "con control total sobre la configuración del servidor."
)
doc.add_paragraph(
    "4. El uso de GitHub facilitó el flujo de trabajo entre desarrollo local y despliegue en la nube."
)
doc.add_paragraph(
    "5. La aplicación es funcional al 100%: registro, login, publicaciones con imágenes, "
    "likes, comentarios y gestión de amistades."
)
doc.add_paragraph(
    "6. Se implementó HTTPS con certificado SSL autofirmado, garantizando comunicación "
    "cifrada entre el cliente y el servidor. Esto permitió habilitar la API getUserMedia "
    "para acceso a la cámara desde el navegador."
)
doc.add_paragraph(
    "7. La cámara WebRTC funciona correctamente en dispositivos móviles y de escritorio, "
    "con soporte para cámara frontal y posterior, vista previa en vivo, y guardia de doble "
    "clic para evitar publicaciones duplicadas."
)
doc.add_paragraph(
    "8. Se corrigió el error de autenticación que retornaba HTTP 500 en lugar de 401, "
    "implementando el manejo correcto de AuthenticationException para APIs stateless."
)

doc.add_paragraph()

# ---- URLs de acceso ----
doc.add_heading("Anexo: URLs de Acceso", level=1)
doc.add_paragraph(f"Frontend Web: https://20.151.96.66/app/")
doc.add_paragraph(f"API Root: https://20.151.96.66/api/")
doc.add_paragraph(f"Nota: HTTP (puerto 80) redirige automáticamente a HTTPS (puerto 443).")
doc.add_paragraph(f"Repositorio GitHub: https://github.com/HYNR24/intercagram")

# Save
doc.save("C:\\Users\\Heyner\\Desktop\\intercagram-main\\Informe_Despliegue_Intercagram_Azure.docx")
print("Informe generado exitosamente.")
